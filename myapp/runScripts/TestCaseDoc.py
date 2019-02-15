# -*- coding:gbk -*-
from __future__ import unicode_literals
from docx import Document
from rest_framework.views import APIView
from ..models import Case, Project
import logging
import time
from django.http import FileResponse
import subprocess
import os


class MakeCases(APIView):
    def post(self, request):
        case_ids = request.data
        global project_name, request_type, request_param, request_url, \
            return_result, case_name, expected_result
        global document
        document = Document()
        # 需要输出的返回信息列表
        records = []
        list_dict = {}
        try:
            # 获取要返回的用例信息
            for number_id in case_ids['ids']:
                project_name = Case.objects.get(isdelete=True,
                                                id=int(
                                                    number_id)).project_name.project_name
                case_name = Case.objects.get(isdelete=True,
                                             id=int(number_id)).case_name
                return_result = Case.objects.get(isdelete=True,
                                                 id=int(
                                                     number_id)).return_result
                request_type = Case.objects.get(isdelete=True,
                                                id=int(number_id)).request_type
                request_param = Case.objects.get(isdelete=True, id=int(
                    number_id)).request_param
                url = Case.objects.get(isdelete=True, id=int(number_id)).url
                permanent_address = Project.objects.get(isdelete=True,
                                                        project_name=Case.objects.get(
                                                            isdelete=True,
                                                            id=int(
                                                                number_id)).project_name).permanent_address
                expected_result = Case.objects.get(isdelete=True, id=int(
                    number_id)).expected_result
                request_url = permanent_address + url
                if request_param == '':
                    request_param = '{}'
                list_dict['request_url'] = request_url
                list_dict['request_type'] = request_type
                list_dict['request_param'] = request_param
                list_dict['expected_result'] = expected_result
                records.append(list_dict)
                document.add_heading(u'接口测试用例', 0)
                document.add_heading('测试用例名称', level=1)
                document.add_heading(case_name, level=1)
                table = document.add_table(rows=1, cols=4)
                hdr_cells = table.rows[0].cells
                # 表格的标题
                hdr_cells[0].text = '请求url'
                hdr_cells[1].text = '请求方式'
                hdr_cells[2].text = '提交参数'
                hdr_cells[3].text = '预期结果'
                row_cells = table.add_row().cells
                # 表格的信息
                row_cells[0].text = list_dict['request_url']
                row_cells[1].text = list_dict['request_type']
                row_cells[2].text = list_dict['request_param']
                row_cells[3].text = list_dict['expected_result']

        except Exception as e:
            logging.info(e)
        word_name = time.strftime("%Y-%m-%d %H:%M:%S",
                                  time.localtime(time.time())) + '.doc'
        document.save(word_name)
        # 获取目标文档的绝对路径
        path = os.path.join(os.path.abspath(os.path.dirname(".")), word_name)
        logging.info("移动目标文档到指定目录")
        subprocess.call(['mv', '-f', path, '/docx'])
        filenames = []
        file = os.listdir('/docx')
        for filename in file:
            if filename.split('.')[1] == 'doc':
                filenames.append(filename.split('.')[0])
        # 获取最新时间的word
        max_doxs = max(filenames)
        name = max_doxs + '.doc'
        logging.info("传输文档内容给前端")
        # 返回word内容给前端
        file = open('/docx/' + name, 'rb')
        response = FileResponse(file)
        response['Content-Type'] = 'application/msword;charset=GB2312'
        response['Content-Disposition'] = 'attachment;filename=' + name
        records = []
        list_dict = {}
        return response
